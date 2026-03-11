"""
Resume Parser Service

This module handles the core logic for parsing resumes in PDF and DOCX formats.
Extracts structured information including name, skills, education, and experience.

Uses heuristic-based extraction methods since we're not using paid APIs.
"""

import pdfplumber
from docx import Document
from typing import Dict, List, Optional
import re
from pathlib import Path

# Optional spaCy import - handle gracefully if not available
try:
    import spacy
    SPACY_AVAILABLE = True
except (ImportError, Exception):
    SPACY_AVAILABLE = False
    spacy = None

from utils.text_cleaner import clean_text, extract_sentences, normalize_skill_name
from .skill_database import (
    SKILL_DATABASE, 
    normalize_skill, 
    is_valid_skill, 
    get_canonical_skill_name
)


_SECTION_END_RE = re.compile(
    r"(?im)^\s*(summary|profile|objective|skills?|technical skills?|projects?|project experience|project|experience|work experience|work history|employment|professional experience|internship|internships|education|certifications?|certificates?|courses?|achievements?|publications?|languages?|interests?|hobbies?)\s*[:\-]?\s*$"
)


def _normalize_line(line: str) -> str:
    line = line.strip()
    # Remove leading bullets / numbering
    line = re.sub(r"^[\s•\-\u2013\u2014\*\u00b7]+", "", line)
    line = re.sub(r"^\(?\d+[\)\.\-]\s*", "", line)
    # Collapse whitespace
    line = re.sub(r"\s+", " ", line).strip()
    return line


def _is_noise_line(line: str) -> bool:
    if not line:
        return True

    low = line.lower().strip()

    # Contact / links / metadata
    if any(k in low for k in ["email", "e-mail", "phone", "mobile", "linkedin", "github", "portfolio", "address", "contact"]):
        return True
    if "@" in line and "." in line:
        return True
    if "http://" in low or "https://" in low or "www." in low:
        return True

    # Section headings (avoid capturing headings as entries)
    if _SECTION_END_RE.match(line):
        return True

    # Very short / generic lines
    if len(line) < 6:
        return True
    if low in {"education", "experience", "work experience", "professional experience", "employment"}:
        return True

    return False


def _extract_section_block(text: str, headings: List[str]) -> str:
    """
    Extract the text block under the first matching heading until the next heading-like line.
    This is intentionally conservative to reduce cross-section bleeding.
    """
    if not text:
        return ""

    # Build start regex: heading must appear on its own line (optionally with ':' or '-')
    heading_re = r"(?im)^\s*(%s)\s*[:\-]?\s*$" % "|".join(re.escape(h) for h in headings)
    m = re.search(heading_re, text)
    if not m:
        return ""

    start = m.end()
    after = text[start:]

    # Find next section heading
    next_m = _SECTION_END_RE.search(after)
    end = next_m.start() if next_m else len(after)
    return after[:end].strip()


def _extract_all_section_blocks(text: str, headings: List[str]) -> str:
    """
    Extract and concatenate blocks for all matching headings, in document order.
    Useful because many resumes split Experience into Internship/Projects/etc.
    """
    if not text:
        return ""

    heading_re = re.compile(
        r"(?im)^\s*(%s)\s*[:\-]?\s*$" % "|".join(re.escape(h) for h in headings)
    )

    blocks: List[str] = []
    for m in heading_re.finditer(text):
        start = m.end()
        after = text[start:]
        next_m = _SECTION_END_RE.search(after)
        end = next_m.start() if next_m else len(after)
        block = after[:end].strip()
        if block:
            blocks.append(block)

    return "\n".join(blocks).strip()


class ResumeParser:
    """
    Main resume parser class that handles file parsing and information extraction.
    """
    
    def __init__(self):
        """
        Initialize the parser with spaCy model and skill keywords.
        """
        # Load spaCy small English model (optional)
        # Note: spaCy may not be compatible with Python 3.14+
        # The parser will work without spaCy using heuristic methods
        self.nlp = None
        if SPACY_AVAILABLE:
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except (OSError, Exception):
                # If model not found or incompatible, we'll work without it
                self.nlp = None
                print("Warning: spaCy not available. Parser will use heuristic methods only.")
        
        # Use comprehensive skill database for accurate extraction
        self.skill_database = SKILL_DATABASE
        
        # Education degree keywords
        self.education_keywords = {
            'bachelor', 'btech', 'b.tech', 'bs', 'b.s', 'bsc', 'b.sc',
            'master', 'mtech', 'm.tech', 'ms', 'm.s', 'msc', 'm.sc', 'mba',
            'phd', 'ph.d', 'doctorate', 'diploma', 'certificate', 'degree'
        }
        
        # Experience indicators
        self.experience_indicators = {
            'worked', 'developed', 'implemented', 'designed', 'created', 'built',
            'managed', 'led', 'collaborated', 'delivered', 'improved', 'optimized',
            'experience', 'years', 'intern', 'internship', 'full-time', 'part-time'
        }
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file using pdfplumber.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text string
        """
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            raise ValueError(f"Error extracting text from PDF: {str(e)}")
        
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file using python-docx.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text string
        """
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")
        
        return text
    
    def extract_name(self, text: str) -> Optional[str]:
        """
        Extract candidate name from resume text.
        
        This is heuristic-based - looks for capitalized words at the beginning
        of the document. Not always accurate, but works for many standard resume formats.
        
        Args:
            text: Cleaned resume text
            
        Returns:
            Candidate name or None if not found
        """
        if not text:
            return None
        
        # Get first few lines (where name typically appears)
        lines = text.split('\n')[:10]
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Look for lines with 2-4 capitalized words (typical name format)
            words = line.split()
            if 2 <= len(words) <= 4:
                # Check if most words start with capital letter
                capitalized = sum(1 for w in words if w and w[0].isupper())
                if capitalized >= len(words) * 0.7:  # At least 70% capitalized
                    # Filter out common resume headers
                    if not any(keyword in line.lower() for keyword in 
                              ['resume', 'cv', 'curriculum', 'vitae', 'email', 'phone', 'address']):
                        return line
        
        return None
    
    def extract_skills(self, text: str) -> List[str]:
        """
        Extract skills from resume text using the skill database.
        
        Only skills that exist in the skill database will be extracted.
        This ensures accuracy and prevents false positives.
        
        Args:
            text: Cleaned resume text
            
        Returns:
            List of extracted skills (canonical names from database)
        """
        if not text:
            return []
        
        found_skills = set()
        text_lower = text.lower()
        
        # First, check for skills in dedicated skill sections
        skill_section_patterns = [
            r'skills?[:\-]?\s*\n(.*?)(?=\n\n|\n[A-Z][a-z]+[:\-]|$)',
            r'technical\s+skills?[:\-]?\s*\n(.*?)(?=\n\n|\n[A-Z][a-z]+[:\-]|$)',
            r'core\s+competencies?[:\-]?\s*\n(.*?)(?=\n\n|\n[A-Z][a-z]+[:\-]|$)',
            r'technologies?[:\-]?\s*\n(.*?)(?=\n\n|\n[A-Z][a-z]+[:\-]|$)',
            r'programming\s+languages?[:\-]?\s*\n(.*?)(?=\n\n|\n[A-Z][a-z]+[:\-]|$)',
        ]
        
        skill_section_text = ""
        for pattern in skill_section_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                skill_section_text += " " + match.group(1)
        
        # Extract skills from skill sections
        if skill_section_text:
            # Split by common delimiters
            items = re.split(r'[,•\-\n|;]', skill_section_text)
            for item in items:
                item = item.strip()
                if 2 <= len(item) <= 50:  # Reasonable skill name length
                    normalized = normalize_skill(item)
                    if normalized in self.skill_database:
                        canonical = get_canonical_skill_name(item)
                        if canonical:
                            found_skills.add(canonical)
        
        # Also search entire text for skill mentions (but prioritize skill sections)
        # Use word boundaries to avoid partial matches
        for skill in self.skill_database:
            # Create pattern with word boundaries
            # Handle multi-word skills
            if ' ' in skill:
                # For multi-word skills, use word boundaries more carefully
                pattern = r'\b' + re.escape(skill) + r'\b'
            else:
                # For single-word skills, use strict word boundaries
                pattern = r'\b' + re.escape(skill) + r'\b'
            
            if re.search(pattern, text_lower):
                canonical = get_canonical_skill_name(skill)
                if canonical:
                    found_skills.add(canonical)
        
        # Return sorted list of unique skills
        return sorted(list(found_skills))
    
    def extract_education(self, text: str) -> List[str]:
        """
        Extract education information from resume with improved filtering.
        
        Looks for degree keywords and extracts surrounding context.
        Filters out noise and invalid entries.
        
        Args:
            text: Cleaned resume text
            
        Returns:
            List of education entries (cleaned and validated)
        """
        if not text:
            return []

        # Prefer strict section extraction to avoid bleeding from other sections.
        edu_block = _extract_section_block(
            text,
            headings=["Education", "Academic Qualifications", "Educational Background", "Academics"],
        )

        # If no explicit section, do a conservative scan of lines.
        candidate_lines = (edu_block.splitlines() if edu_block else text.splitlines())

        # Merge wrapped lines: if a line looks like continuation, append to previous.
        merged: List[str] = []
        for raw in candidate_lines:
            line = _normalize_line(raw)
            if _is_noise_line(line):
                continue

            # Skip obvious non-education headings inside the block
            if any(h in line.lower() for h in ["skills", "experience", "projects", "certification", "summary"]):
                continue

            if merged and (raw.startswith(" ") or raw.startswith("\t") or (line and line[0].islower())):
                merged[-1] = (merged[-1] + " " + line).strip()
            else:
                merged.append(line)

        # Keep only lines that actually look like education entries.
        entries: List[str] = []
        for line in merged:
            low = line.lower()

            # Must contain a degree/education keyword OR a common institute indicator.
            has_degree = any(k in low for k in self.education_keywords)
            has_institute = any(k in low for k in ["university", "college", "institute", "school"])
            has_year = bool(re.search(r"\b(19|20)\d{2}\b", line))
            has_cgpa = bool(re.search(r"\b(cgpa|gpa|grade|percentage|%)\b", low))

            if not (has_degree or has_institute):
                continue

            # Drop pure scores without context
            if has_cgpa and not (has_degree or has_institute):
                continue

            # Basic length sanity
            if not (12 <= len(line) <= 180):
                continue

            # Remove duplicated whitespace and trailing punctuation noise
            cleaned = re.sub(r"\s+", " ", line).strip(" -•")

            # Avoid duplicates
            if not any(e.lower() == cleaned.lower() for e in entries):
                entries.append(cleaned)

        # Prefer lines with years (usually more complete)
        entries.sort(key=lambda s: (0 if re.search(r"\b(19|20)\d{2}\b", s) else 1, -len(s)))
        return entries[:5]
    
    def extract_experience(self, text: str) -> List[str]:
        """
        Extract work experience information from resume with improved filtering.
        
        Looks for sentences that contain experience-related keywords.
        Filters out noise and invalid entries.
        
        Args:
            text: Cleaned resume text
            
        Returns:
            List of experience descriptions (cleaned and validated)
        """
        if not text:
            return []

        # Treat internships as experience, and also include "Projects" style sections
        # because many student resumes list internships under "Internships" and work under "Work History".
        exp_block = _extract_all_section_blocks(
            text,
            headings=[
                "Work Experience",
                "Professional Experience",
                "Work History",
                "Experience",
                "Employment",
                "Internship",
                "Internships",
                "Project Experience",
                "Projects",
            ],
        )

        # If there is no explicit section, do NOT do a broad sentence scrape (it causes noise).
        # Instead do a conservative line scan looking for date ranges + role/company cues.
        candidate_lines = exp_block.splitlines() if exp_block else text.splitlines()

        # Helpers for recognizing experience lines
        date_range_re = re.compile(
            r"(?i)\b((19|20)\d{2}|jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)\b.*\b((19|20)\d{2}|present|current)\b"
        )
        role_cues = [
            # Internship-related (explicitly treated as experience)
            "intern", "internship", "trainee",
            # Common roles
            "engineer", "developer", "analyst", "consultant", "designer",
            "manager", "lead", "assistant", "associate", "tester", "administrator",
        ]
        company_cues = ["pvt", "ltd", "inc", "llc", "company", "technologies", "solutions", "labs", "systems"]
        action_verbs = {"developed", "built", "implemented", "designed", "created", "optimized", "led", "managed", "delivered"}

        merged: List[str] = []
        for raw in candidate_lines:
            line = _normalize_line(raw)
            if _is_noise_line(line):
                continue

            low = line.lower()

            # Skip clearly-not-experience sections within block
            # NOTE: do NOT skip "projects" here because we intentionally include Project sections as experience-like entries.
            if any(h in low for h in ["education", "skills", "certification", "summary"]):
                continue

            # Merge wrapped bullets / continuation lines
            if merged and (raw.startswith(" ") or raw.startswith("\t") or low.startswith(("•", "-", "*")) or (line and line[0].islower())):
                merged[-1] = (merged[-1] + " " + line).strip()
            else:
                merged.append(line)

        # Now filter to meaningful experience entries:
        # - either a role/company header line (often includes date range), OR
        # - a bullet achievement line with an action verb.
        entries: List[str] = []
        for line in merged:
            low = line.lower()

            has_date = bool(date_range_re.search(line)) or bool(re.search(r"\b(19|20)\d{2}\b", line))
            has_role = any(cue in low for cue in role_cues)
            has_company = any(cue in low for cue in company_cues) or (" - " in line) or (" | " in line)
            has_action = any(v in low.split() for v in action_verbs) or any(v in low for v in action_verbs)

            # Reject "Responsibilities:"-style headings
            if low.rstrip(":") in {"responsibilities", "key responsibilities", "projects", "highlights"}:
                continue

            # Avoid lines that are just dates/locations
            if re.fullmatch(r"[\d\s\-\/|,]+", line):
                continue

            # Keep strong signals
            if (has_role and (has_company or has_date)) or (has_action and len(line) >= 35):
                cleaned = re.sub(r"\s+", " ", line).strip(" -•")
                if 20 <= len(cleaned) <= 220 and not any(e.lower() == cleaned.lower() for e in entries):
                    entries.append(cleaned)

        # Prefer header-like lines with dates first, then longer achievement bullets
        entries.sort(
            key=lambda s: (
                0 if date_range_re.search(s) else 1,
                0 if any(cue in s.lower() for cue in role_cues) else 1,
                -len(s),
            )
        )
        return entries[:10]
    
    def parse_resume(self, file_path: str, file_type: str) -> Dict:
        """
        Main method to parse a resume file and extract structured information.
        
        Args:
            file_path: Path to the resume file
            file_type: Type of file ('pdf' or 'docx')
            
        Returns:
            Dictionary containing parsed resume data:
            {
                'name': str or None,
                'skills': List[str],
                'education': List[str],
                'experience': List[str]
            }
        """
        # Extract raw text based on file type
        if file_type.lower() == 'pdf':
            raw_text = self.extract_text_from_pdf(file_path)
        elif file_type.lower() == 'docx':
            raw_text = self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Clean the text
        cleaned_text = clean_text(raw_text)
        
        if not cleaned_text:
            raise ValueError("No text could be extracted from the resume file")
        
        # Extract structured information
        name = self.extract_name(cleaned_text)
        skills = self.extract_skills(cleaned_text)
        education = self.extract_education(cleaned_text)
        experience = self.extract_experience(cleaned_text)
        
        # Final validation and cleaning
        result = {
            'name': name if name and len(name.strip()) > 0 else None,
            'skills': skills if skills else [],
            'education': education if education else [],
            'experience': experience if experience else []
        }
        
        return result
